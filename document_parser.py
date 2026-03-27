import PyPDF2
from docx import Document
import json
import os
import base64
import time
import io

# PyMuPDF (fitz) — best PDF text extraction + page-to-image rendering
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False


def extract_text_from_pdf_fitz(file_path):
    """Primary PDF extraction using PyMuPDF (fitz) — best quality."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            page_text = page.get_text("text")
            if page_text:
                text += page_text + "\n"
        doc.close()
    except Exception as e:
        print(f"PyMuPDF (fitz) error: {e}")
    return text.strip()


def extract_text_from_pdf_pypdf2(file_path):
    """Fallback PDF extraction using PyPDF2."""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    text += extracted + "\n"
    except Exception as e:
        print(f"PyPDF2 error: {e}")
    return text.strip()


GEMINI_MODELS = ['gemini-2.0-flash', 'gemini-1.5-flash']

OCR_PROMPT = (
    "Extract ALL the text from this exam/answer sheet image. "
    "Return ONLY the raw text content exactly as written, preserving "
    "the structure (questions, answers, numbering). "
    "Do NOT add any commentary or interpretation. "
    "If there are handwritten answers, transcribe them as accurately as possible."
)


def _call_gemini_vision(client, model, img_bytes):
    """Call Gemini vision API for a single page image."""
    from google.genai import types
    response = client.models.generate_content(
        model=model,
        contents=[
            types.Content(
                parts=[
                    types.Part(
                        inline_data=types.Blob(
                            mime_type='image/png',
                            data=img_bytes
                        )
                    ),
                    types.Part(text=OCR_PROMPT)
                ]
            )
        ]
    )
    return response.text.strip()


def extract_text_from_scanned_pdf_via_gemini(file_path):
    """
    For scanned/image PDFs: render pages as images → Gemini Vision OCR.
    Tries multiple models (2.0-flash → 1.5-flash) to handle quota limits.
    """
    if not HAS_FITZ:
        return ""
    
    doc = None
    try:
        from google import genai
        from dotenv import load_dotenv

        load_dotenv()
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            print("No GEMINI_API_KEY found for OCR fallback.")
            return ""

        client = genai.Client(api_key=api_key)
        
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        # Render all pages to images first, then close PDF to release file lock
        page_images = []
        for page_num in range(total_pages):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            page_images.append(pix.tobytes("png"))
        
        doc.close()
        doc = None
        
        print(f"  Rendered {total_pages} pages. Starting OCR...")
        
        all_text = ""
        current_model_idx = 0
        
        for page_num, img_bytes in enumerate(page_images):
            success = False
            
            while current_model_idx < len(GEMINI_MODELS):
                model = GEMINI_MODELS[current_model_idx]
                print(f"  OCR page {page_num + 1}/{total_pages} using {model}...")
                
                try:
                    page_text = _call_gemini_vision(client, model, img_bytes)
                    if page_text:
                        all_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    success = True
                    break
                    
                except Exception as api_err:
                    err_str = str(api_err)
                    if '429' in err_str or 'RESOURCE_EXHAUSTED' in err_str:
                        if 'PerDay' in err_str or 'limit: 0' in err_str:
                            # Daily quota exhausted — switch model immediately
                            print(f"  ❌ {model} daily quota exhausted. Switching model...")
                            current_model_idx += 1
                        else:
                            # Per-minute limit — wait briefly and retry same model
                            print(f"  ⏳ {model} rate limited. Waiting 10s...")
                            time.sleep(10)
                    else:
                        print(f"  ❌ API error on page {page_num + 1}: {api_err}")
                        current_model_idx += 1
            
            if not success and current_model_idx >= len(GEMINI_MODELS):
                print(f"  ⚠️ All models exhausted. OCR'd {page_num}/{total_pages} pages.")
                break
        
        return all_text.strip()
        
    except Exception as e:
        print(f"Gemini Vision OCR error: {e}")
        return ""
    finally:
        if doc is not None:
            try:
                doc.close()
            except:
                pass


def extract_text_from_file(file_path, filename):
    """Reads raw text from PDF, DOCX, TXT, MD, or JSON files.
    For scanned/image PDFs: renders pages as images → sends to Gemini Vision for OCR."""
    text = ""
    ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        if ext in ['txt', 'md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                
        elif ext == 'json':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    text = json.dumps(data, indent=2)
            except json.JSONDecodeError as e:
                return f"[FORMAT_ERROR: Invalid JSON Syntax. {str(e)}]"
                
        elif ext == 'docx':
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables (common in exam papers)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            
            text = "\n".join(paragraphs)
            
        elif ext == 'pdf':
            # Method 1: PyMuPDF text extraction (works for text-based PDFs)
            if HAS_FITZ:
                text = extract_text_from_pdf_fitz(file_path)
            
            # Method 2: PyPDF2 fallback
            if len(text.strip()) < 10:
                print("Fitz extracted very little, trying PyPDF2 fallback...")
                fallback = extract_text_from_pdf_pypdf2(file_path)
                if len(fallback) > len(text):
                    text = fallback
            
            # Method 3: SCANNED PDF — Use Gemini Vision OCR
            if len(text.strip()) < 10:
                print("PDF appears to be scanned/image-based. Using Gemini Vision OCR...")
                ocr_text = extract_text_from_scanned_pdf_via_gemini(file_path)
                if len(ocr_text) > len(text):
                    text = ocr_text
            
            # If ALL methods fail
            if len(text.strip()) < 10:
                return "[FORMAT_ERROR: Could not extract text from this PDF even with AI OCR. The file may be corrupted or completely blank.]"
                    
    except Exception as e:
        text = f"[ERROR READING FILE: {filename} - {str(e)}]"

    return text